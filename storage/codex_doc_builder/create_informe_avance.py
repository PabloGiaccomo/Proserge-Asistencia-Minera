from copy import deepcopy
from pathlib import Path
import tempfile
import time
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


TEMPLATE = Path(r"C:\Users\giacc\Downloads\ActaGestiondeCambios20-05-2026.docx")
OUTPUT = Path(r"C:\proserge_app\storage\codex_outputs\Informe_Avance_PROSERGE_19-06-2026_LISTO.docx")
RAW_OUTPUT = Path(r"C:\proserge_app\storage\codex_outputs\_tmp_Informe_Avance_PROSERGE.docx")

BLUE = "0B5FA5"
DARK_BLUE = "17365D"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9D9D9"
WHITE = "FFFFFF"
BLACK = RGBColor(0, 0, 0)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_font(run, size=10.5, bold=False, italic=False, color=BLACK, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def paragraph(text="", size=10.5, bold=False, italic=False, color=BLACK, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(text, level=1):
    sizes = {1: 12.5, 2: 11.5, 3: 10.8}
    before = {1: 12, 2: 8, 3: 6}
    p = paragraph(text, size=sizes.get(level, 11), bold=True, color=RGBColor(0x17, 0x36, 0x5D), before=before.get(level, 8), after=6)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run("• ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="A6A6A6", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def cell_text(cell, text, bold=False, size=9.2, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    set_cell_borders(cell)


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Cm(width)


def fixed_table(rows, cols, widths=None, header=True):
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    if widths:
        set_table_widths(table, widths)
    mark_header_row(table.rows[0])
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell)
            set_cell_margins(cell)
            if header and i == 0:
                set_cell_shading(cell, BLUE)
    return table


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def patch_header_image_alt(source_path: Path, target_path: Path):
    ns = {
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    with zipfile.ZipFile(source_path, "r") as zin, zipfile.ZipFile(target_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                root = ET.fromstring(data)
                for doc_pr in root.findall(".//wp:docPr", ns):
                    doc_pr.set("title", "Membrete CDPERU")
                    doc_pr.set("descr", "Membrete institucional de CDPERU con datos de contacto.")
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(item, data)


def fill_row(row, values, header=False, label_cols=None, alignments=None, size=9.2):
    label_cols = label_cols or set()
    alignments = alignments or {}
    for i, value in enumerate(values):
        if header:
            set_cell_shading(row.cells[i], BLUE)
            cell_text(row.cells[i], value, bold=True, size=size, color=RGBColor(255, 255, 255), align=alignments.get(i, WD_ALIGN_PARAGRAPH.CENTER))
        else:
            if i in label_cols:
                set_cell_shading(row.cells[i], LIGHT_GRAY)
            cell_text(row.cells[i], value, bold=i in label_cols, size=size, align=alignments.get(i, WD_ALIGN_PARAGRAPH.LEFT))


def add_identification_table():
    table = fixed_table(6, 4, widths=[2.7, 4.5, 2.5, 4.8], header=False)
    rows = [
        ("Proyecto:", "PROSERGE - Sistema de Gestión Operativa de Personal en Minería", "Código informe:", "AV-PROSERGE-002"),
        ("Versión:", "V1.0", "Fecha:", "19 de junio de 2026"),
        ("Elaborado por:", "Equipo de Desarrollo PROSERGE / CDPERU", "Dirigido a:", "Gerencia y jefaturas de área"),
        ("Periodo informado:", "Avance acumulado a la fecha", "Estado general:", "En desarrollo - avance condicionado por validación operativa"),
        ("Alcance principal:", "Personal, contratos, fichas, documentos, habilitación minera, bienestar, RQ y herramientas", "Documento base:", "Acta de Gestión de Cambios 20-05-2026"),
        ("Objetivo:", "Documentar avance, pendientes, retrasos y acciones requeridas", "Prioridad:", "Alta - requiere respuesta de áreas usuarias"),
    ]
    for row, values in zip(table.rows, rows):
        fill_row(row, values, label_cols={0, 2}, size=8.9)


def add_status_table():
    heading("4. AVANCE POR MÓDULO Y SITUACIÓN ACTUAL")
    table = fixed_table(8, 6, widths=[0.8, 2.1, 1.7, 4.0, 4.0, 2.0], header=True)
    fill_row(table.rows[0], ["N°", "Módulo", "Estado", "Avances documentados", "Pendiente / bloqueo actual", "Responsable de avance"], header=True, size=8.2)
    rows = [
        ("1", "Personal", "Avance alto", "Gran parte del módulo se encuentra desarrollado: registro, ficha, documentos, revisión, personal nuevo y regularización de datos solicitados por RR.HH.", "Queda pendiente automatizar la descarga consolidada en PDF y cerrar observaciones finas solicitadas por RR.HH.", "Desarrollo + RR.HH."),
        ("2", "Contratos", "Implementado en ajuste", "Se incorporó el flujo de contratos a pedido de Elida, incluyendo datos contractuales y control de contrato firmado.", "Falta que el área se acostumbre al uso del programa y revise correctamente las fichas antes de aprobarlas.", "RR.HH."),
        ("3", "Personal antiguo", "Bloqueado por insumo", "Se solicitó un formato para registrar y validar a todo el personal que ya estaba en la empresa.", "Han transcurrido aproximadamente 3 semanas sin recepción del formato/listado requerido, impidiendo revisión masiva y depuración de errores.", "Jefaturas / RR.HH."),
        ("4", "Habilitación minera", "Pendiente de arranque operativo", "El flujo está identificado como parte crítica del sistema: trabajador, mina, requisitos, exámenes, intentos, resultado y estado.", "A pesar de haberse solicitado, aún no se inicia la revisión/avance operativo con todo el personal para detectar errores y mejoras.", "Operaciones / SSOMA"),
        ("5", "Bienestar", "Sin validación funcional", "El módulo contempla bloqueos y restricciones de trabajadores.", "Aún no se registran bloqueos reales para validar el comportamiento del módulo y su impacto operativo.", "Bienestar / SSOMA"),
        ("6", "RQ Mina / RQ Proserge", "Aprobado para continuar", "Los planners dieron visto bueno para avanzar con RQ Proserge a partir de la lógica de RQ Mina.", "Requiere validación de flujo real y comunicación clara de cambios por parte de responsables.", "Planners / Operaciones"),
        ("7", "Herramientas y logística", "En ampliación de alcance", "Se incorporó el pedido de herramientas, considerando costos y gastos coordinados con Logística.", "Debe cerrarse la definición de responsables, costos, validaciones y reportes esperados.", "Logística / Desarrollo"),
    ]
    for row, values in zip(table.rows[1:], rows):
        fill_row(row, values, alignments={0: WD_ALIGN_PARAGRAPH.CENTER, 2: WD_ALIGN_PARAGRAPH.CENTER}, size=7.05)


def add_delay_table():
    heading("6. RETRASOS, CAUSAS Y EFECTO EN EL PROYECTO")
    paragraph(
        "Los retrasos identificados no se originan en una única funcionalidad, sino en la falta de comunicación formal, "
        "la demora en la entrega de información por área y el bajo uso inicial del programa por parte de los responsables operativos. "
        "Esto obliga a invertir más tiempo en averiguar procesos que deberían ser comunicados con claridad por cada jefatura."
    )
    table = fixed_table(6, 4, widths=[3.0, 4.3, 4.3, 2.0], header=True)
    fill_row(table.rows[0], ["Factor", "Situación registrada", "Impacto", "Nivel"], header=True, size=8.4)
    rows = [
        ("Comunicación por áreas", "Las jefaturas no entregan con claridad o a tiempo los procesos, necesidades y responsables.", "El desarrollo dedica tiempo adicional a investigar procesos, corregir supuestos y rehacer ajustes.", "Alto"),
        ("Formato de personal existente", "Se pidió un formato/listado para registrar al personal ya existente y aún no se recibe luego de aproximadamente 3 semanas.", "No se puede validar de forma masiva al personal antiguo ni detectar duplicados, errores o mejoras.", "Crítico"),
        ("Adopción del sistema", "Las áreas todavía no empiezan a trabajar de forma constante sobre el programa.", "La retroalimentación llega tarde y los errores de uso aparecen después de implementar.", "Alto"),
        ("Revisión de fichas", "RR.HH. debe revisar las fichas correctamente para identificar errores de carga cometidos por trabajadores.", "Si no se revisa a tiempo, se aprueban datos incompletos o se acumulan correcciones.", "Medio/Alto"),
        ("Habilitación minera y bienestar", "No se ha iniciado la carga/revisión operativa esperada ni los bloqueos reales.", "No se puede validar el funcionamiento completo de esos módulos con casos reales.", "Alto"),
    ]
    for row, values in zip(table.rows[1:], rows):
        fill_row(row, values, alignments={3: WD_ALIGN_PARAGRAPH.CENTER}, size=8.0)


def add_action_table():
    heading("8. ACCIONES REQUERIDAS PARA RECUPERAR RITMO")
    table = fixed_table(7, 5, widths=[0.8, 4.6, 2.7, 2.7, 3.8], header=True)
    fill_row(table.rows[0], ["N°", "Acción requerida", "Responsable sugerido", "Plazo sugerido", "Resultado esperado"], header=True, size=8.2)
    rows = [
        ("1", "Entregar el formato/listado del personal antiguo y vigente de la empresa.", "RR.HH. / jefaturas", "Inmediato", "Permitir carga, depuración y revisión integral del personal."),
        ("2", "Iniciar revisión de habilitación minera con trabajadores reales.", "Operaciones / SSOMA", "Semana actual", "Detectar errores, brechas y mejoras antes de cierre funcional."),
        ("3", "Registrar bloqueos reales de trabajadores en Bienestar.", "Bienestar", "Semana actual", "Validar impacto de bloqueos y restricciones en disponibilidad."),
        ("4", "Revisar fichas del personal nuevo antes de aprobarlas.", "RR.HH.", "Continuo", "Evitar errores por carga incorrecta de trabajadores."),
        ("5", "Cerrar criterios de costos y gastos para herramientas con Logística.", "Logística", "Corto plazo", "Definir reportes y trazabilidad económica."),
        ("6", "Formalizar reuniones cortas por área con acuerdos escritos.", "Gerencia / áreas", "Semanal", "Reducir reprocesos por falta de comunicación."),
    ]
    for row, values in zip(table.rows[1:], rows):
        fill_row(row, values, alignments={0: WD_ALIGN_PARAGRAPH.CENTER, 3: WD_ALIGN_PARAGRAPH.CENTER}, size=8.0)


def add_roadmap_table():
    heading("7. CRONOGRAMA DE SEGUIMIENTO OPERATIVO")
    paragraph("La siguiente vista resume el foco inmediato de trabajo y validación por módulo. No reemplaza el Gantt general; funciona como control de avance y desbloqueo.")
    table = fixed_table(7, 6, widths=[3.7, 2.0, 2.0, 2.0, 2.0, 3.0], header=True)
    fill_row(table.rows[0], ["Actividad", "Actual", "Semana +1", "Semana +2", "Semana +3", "Criterio de cierre"], header=True, size=8.0)
    rows = [
        ("Personal nuevo y ajustes RR.HH.", "●", "●", "●", "", "Fichas revisadas y observaciones cerradas"),
        ("PDF automático de ficha/descarga", "●", "●", "", "", "Descarga generada desde el sistema"),
        ("Personal antiguo - formato y validación", "Pendiente insumo", "●", "●", "●", "Listado recibido, depurado y sin duplicados"),
        ("Habilitación minera - revisión real", "Pendiente", "●", "●", "●", "Trabajador-mina-exámenes validados"),
        ("Bienestar - bloqueos", "Pendiente", "●", "●", "", "Bloqueos reales registrados y probados"),
        ("RQ Proserge / herramientas / costos", "●", "●", "●", "●", "Validación operativa y costos confirmados"),
    ]
    for row, values in zip(table.rows[1:], rows):
        fill_row(row, values, alignments={1: WD_ALIGN_PARAGRAPH.CENTER, 2: WD_ALIGN_PARAGRAPH.CENTER, 3: WD_ALIGN_PARAGRAPH.CENTER, 4: WD_ALIGN_PARAGRAPH.CENTER}, size=7.8)
        for i in range(1, 5):
            val = values[i]
            if val == "●":
                set_cell_shading(row.cells[i], LIGHT_BLUE)
            elif "Pendiente" in val:
                set_cell_shading(row.cells[i], "FCE4D6")


def add_footer_signature():
    heading("9. OBSERVACIONES FINALES")
    paragraph(
        "El proyecto mantiene avance técnico importante, especialmente en Personal, fichas, documentos y contratos. "
        "Sin embargo, el cierre efectivo depende de que las áreas usuarias entreguen información, usen el sistema, revisen datos reales "
        "y comuniquen con precisión sus necesidades. Cada retraso en validación operativa posterga la detección de errores y la estabilización del programa."
    )
    paragraph(
        "Se recomienda que Gerencia solicite a cada responsable de área una respuesta formal con entregables, fechas y responsables, "
        "especialmente para personal antiguo, habilitación minera, bienestar y costos/logística."
    )
    paragraph("Elaborado para seguimiento de avance y toma de decisiones.", size=9.5, italic=True, color=RGBColor(85, 85, 85), before=8)


doc = Document(TEMPLATE)
clear_body(doc)

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(5.0)
section.bottom_margin = Cm(5.0)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

paragraph("Informe de Avance", size=12, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
paragraph("INFORME DE AVANCE DEL PROYECTO PROSERGE", size=15.5, bold=True, color=RGBColor(0x17, 0x36, 0x5D), align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
paragraph("Sistema PROSERGE - Gestión Operativa de Personal en Minería", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

heading("1. IDENTIFICACIÓN DEL INFORME")
add_identification_table()

heading("2. ANTECEDENTES Y CONTEXTO")
paragraph(
    "El proyecto PROSERGE continúa en desarrollo como sistema corporativo para ordenar procesos de Personal, contratos, documentos, "
    "habilitación minera, requerimientos, herramientas, asistencia y áreas relacionadas. El acta de gestión de cambios del 20 de mayo de 2026 "
    "ya dejó constancia de modificaciones de alcance y ajustes al cronograma. El presente informe actualiza el avance real, los pendientes y las dependencias operativas vigentes."
)
bullet("El módulo de Personal concentra el mayor avance y ya cubre gran parte de los flujos necesarios para personal nuevo, revisión de fichas, documentos y contratos.")
bullet("El sistema incorporó contratos a pedido de Elida; el siguiente reto no es solo técnico, sino de adopción y revisión constante por parte de RR.HH.")
bullet("Habilitación minera, Bienestar, personal antiguo y validaciones masivas siguen dependiendo de información y uso real por parte de las áreas responsables.")
bullet("Los retrasos se explican principalmente por falta de comunicación formal, demora en entregables por área y ausencia de validación operativa oportuna.")

heading("3. ALCANCE EXISTENTE Y BASE DEL PROYECTO")
paragraph("Antes del presente informe, el proyecto ya contemplaba una ruta integral de operación minera y administrativa:")
bullet("Personal: registro, ficha pública, documentos, revisión, observación, aprobación, estados laborales y descargas documentales.")
bullet("Contratos: datos contractuales, contrato firmado, renovaciones, historial y control para evitar activar trabajadores sin contrato vigente firmado.")
bullet("Habilitación minera: relación trabajador-mina, requisitos/exámenes, intentos, resultados, convalidaciones y estados de habilitación.")
bullet("RQ Mina y RQ Proserge: requerimientos operativos, planificación, asignación de personal y continuidad hacia Man Power y asistencia.")
bullet("Herramientas y logística: requerimientos de herramientas por parada, ahora ampliados para contemplar costos y gastos.")
bullet("Bienestar: bloqueos, restricciones y condiciones que deben probarse con casos reales.")

add_status_table()

heading("5. AVANCE NARRATIVO POR ÁREA")
heading("5.1 Personal y RR.HH.", level=2)
paragraph(
    "Personal es el módulo más avanzado. Ya se viene trabajando con personal nuevo y se están corrigiendo los detalles que RR.HH. sigue solicitando. "
    "La prioridad actual es cerrar la automatización de descarga en PDF y reforzar el proceso de revisión de fichas, porque los trabajadores pueden cometer errores al cargar información y documentos."
)
heading("5.2 Contratos", level=2)
paragraph(
    "El flujo de contratos fue agregado a pedido de Elida. El sistema ya considera el contrato como parte central del ciclo laboral; falta consolidar el hábito de uso, "
    "subida/revisión de contratos firmados y revisión correcta de fichas para no aprobar información incompleta."
)
heading("5.3 Habilitación minera", level=2)
paragraph(
    "La habilitación minera debe iniciar con validación real de trabajadores, minas, requisitos, exámenes e intentos. Aunque ya se mencionó la necesidad de empezar, "
    "aún no se ha iniciado el avance operativo con todo el personal para verificar errores, identificar mejoras y ordenar la información."
)
heading("5.4 Bienestar", level=2)
paragraph(
    "El módulo de Bienestar continúa sin registrar bloqueos reales de trabajadores. Sin esos casos no se puede comprobar el comportamiento funcional de restricciones, bloqueos y disponibilidad."
)
heading("5.5 RQ Mina, RQ Proserge, herramientas y logística", level=2)
paragraph(
    "Los planners dieron el visto bueno para avanzar con RQ Proserge a partir de RQ Mina. También se avanzó con el pedido de herramientas, incorporando el análisis de costos y gastos junto con Logística. "
    "Queda pendiente cerrar responsables, criterios de validación y reportes esperados."
)

add_delay_table()
add_roadmap_table()
add_action_table()
add_footer_signature()

doc.core_properties.title = "Informe de Avance del Proyecto PROSERGE"
doc.core_properties.subject = "Avance, pendientes, retrasos y acciones requeridas"
doc.core_properties.author = "CDPERU / Equipo de Desarrollo PROSERGE"
doc.core_properties.comments = "Documento generado como informe de avance del proyecto PROSERGE."

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(RAW_OUTPUT)
patch_header_image_alt(RAW_OUTPUT, OUTPUT)
try:
    RAW_OUTPUT.unlink()
except OSError:
    pass
print(OUTPUT)
