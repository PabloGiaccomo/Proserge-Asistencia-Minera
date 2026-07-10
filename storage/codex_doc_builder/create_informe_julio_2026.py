from pathlib import Path
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


TEMPLATE = Path(r"C:\Users\giacc\Downloads\ActaGestiondeCambios20-05-2026.docx")
OUTPUT = Path(r"C:\proserge_app\storage\codex_outputs\Informe_Mensual_Labores_Julio_2026_PROSERGE.docx")
RAW_OUTPUT = Path(r"C:\proserge_app\storage\codex_outputs\_tmp_Informe_Mensual_Julio_2026.docx")

BLUE = "0B5FA5"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
BLACK = RGBColor(0, 0, 0)
DARK_BLUE = RGBColor(0x17, 0x36, 0x5D)
MUTED = RGBColor(85, 85, 85)


def clear_body(doc: Document) -> None:
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_font(run, size=10.3, bold=False, italic=False, color=BLACK, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def paragraph(doc, text="", size=10.3, bold=False, italic=False, color=BLACK, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 12.5, 2: 11.3, 3: 10.8}
    before = {1: 11, 2: 7, 3: 5}
    return paragraph(
        doc,
        text,
        size=sizes.get(level, 11),
        bold=True,
        color=DARK_BLUE,
        before=before.get(level, 8),
        after=5,
    )


def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run("- ")
    set_font(r, size=10.3, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.3)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="A6A6A6", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_text(cell, text, bold=False, size=8.4, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    set_cell_borders(cell)


def fixed_table(doc, rows, cols, widths=None, header=True):
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    mark_header_row(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell)
            set_cell_margins(cell)
            if header and row_idx == 0:
                set_cell_shading(cell, BLUE)
    return table


def fill_row(row, values, header=False, label_cols=None, alignments=None, size=8.4):
    label_cols = label_cols or set()
    alignments = alignments or {}
    for i, value in enumerate(values):
        if header:
            set_cell_shading(row.cells[i], BLUE)
            cell_text(row.cells[i], value, bold=True, size=size, color=RGBColor(255, 255, 255), align=alignments.get(i, WD_ALIGN_PARAGRAPH.CENTER))
        else:
            if i in label_cols:
                set_cell_shading(row.cells[i], LIGHT_GRAY)
            cell_text(row.cells[i], value, bold=i in label_cols, size=size, align=alignments.get(i, WD_ALIGN_PARAGRAPH.LEFT))


def add_identification(doc):
    table = fixed_table(doc, 6, 4, widths=[2.6, 4.7, 2.5, 4.8], header=False)
    rows = [
        ("Documento:", "Informe mensual de labores", "Código:", "INF-JUL-PROSERGE-001"),
        ("Periodo:", "01 al 30 de julio de 2026", "Fecha de elaboración:", "01 de julio de 2026"),
        ("Proyecto:", "PROSERGE - Sistema de Gestión Operativa", "Área:", "Sistemas / Soporte / Desarrollo"),
        ("Cliente interno:", "PROSERGE", "Formato:", "CDPERU"),
        ("Alcance:", "Soporte TI, nube, correos, Excel, equipos y sistema Proserge", "Estado:", "Informe mensual preparado"),
        ("Nota:", "Las actividades se redactan de forma formal para presentación interna.", "Confidencialidad:", "Uso interno"),
    ]
    for row, values in zip(table.rows, rows):
        fill_row(row, values, label_cols={0, 2}, size=8.3)


def add_labor_table(doc):
    heading(doc, "3. LABORES REALIZADAS EN EL MES")
    table = fixed_table(doc, 9, 5, widths=[0.8, 3.2, 3.1, 5.0, 2.5], header=True)
    fill_row(table.rows[0], ["N°", "Actividad", "Área relacionada", "Detalle de avance", "Estado"], header=True, size=7.9)
    rows = [
        ("1", "Reuniones de sistema", "Gestión / áreas usuarias", "Participación en reuniones de coordinación para revisar necesidades del sistema, levantar observaciones y ordenar prioridades de avance.", "Realizado"),
        ("2", "Seguimiento de nube Nextcloud", "Infraestructura TI", "Seguimiento del servicio de nube, validación de acceso, sincronización y continuidad operativa para archivos compartidos.", "Realizado"),
        ("3", "Configuración de correos", "Soporte TI", "Configuración y revisión de cuentas de correo en celulares y computadoras, incluyendo pruebas de envío, recepción y acceso de usuarios.", "Realizado"),
        ("4", "Soporte Excel a RR.HH.", "RR.HH.", "Corrección de una fórmula extensa de Excel para Carmen, asistente de RR.HH., con validación funcional para su uso operativo.", "Realizado"),
        ("5", "Avance de logística en Proserge", "Logística / Operaciones", "Avance funcional para incorporar logística al sistema: herramientas, EPP, seguimiento de entregas, costos/gastos y transporte asociado a RQ Mina.", "En avance"),
        ("6", "Soporte de suite Office", "Soporte TI", "Revisión y habilitación operativa de Office en computadoras de oficina, redactado como soporte de activación/licenciamiento y disponibilidad de herramientas ofimáticas.", "Realizado"),
        ("7", "Revisión de impresoras", "Soporte TI", "Revisión de impresoras de oficina, conectividad, funcionamiento y disponibilidad para usuarios.", "Realizado"),
        ("8", "Traspaso de información", "Soporte TI / respaldo", "Migración de información desde un disco duro liviano hacia un disco de mayor capacidad, con objetivo de resguardo y continuidad de archivos.", "Realizado"),
    ]
    for row, values in zip(table.rows[1:], rows):
        fill_row(row, values, alignments={0: WD_ALIGN_PARAGRAPH.CENTER, 4: WD_ALIGN_PARAGRAPH.CENTER}, size=7.55)


def add_system_changes(doc):
    heading(doc, "4. CAMBIOS Y AVANCES DEL SISTEMA PROSERGE")
    paragraph(
        doc,
        "La revisión del repositorio muestra avances relevantes en módulos operativos y de logística. No se detectaron commits registrados entre el 01 y el 30 de julio de 2026, por lo que esta sección resume cambios presentes en el entorno de trabajo y avances funcionales preparados para el periodo informado.",
        size=9.8,
        color=MUTED,
    )
    table = fixed_table(doc, 8, 5, widths=[0.8, 3.1, 4.0, 4.8, 2.0], header=True)
    fill_row(table.rows[0], ["N°", "Módulo", "Cambio incorporado", "Impacto operativo", "Estado"], header=True, size=7.7)
    rows = [
        ("1", "Logística EPP", "Se incorporó estructura para catálogo de EPP, registro de entregas, cambios/devoluciones, vida útil, vencimiento calendario y uso efectivo por trabajador.", "Permite controlar EPP entregado, relacionarlo con personal y calcular uso por paradas o asignaciones.", "En avance"),
        ("2", "RQ Mina - transporte", "Se amplió el seguimiento logístico de transporte con origen, placas asignadas, fechas, días de uso, estado logístico, recepción e incidencias.", "Mejora trazabilidad de transporte requerido para operaciones mineras y permite registrar eventos/cambios.", "En avance"),
        ("3", "Herramientas por parada", "Se agregaron fases de entrega y recepción, cantidades entregadas/recibidas, incidencias durante parada y comentarios por cambios previos.", "Ayuda a comparar lo solicitado, entregado y recibido, reduciendo pérdida de control en logística de parada.", "En avance"),
        ("4", "Personal e ingresos", "Se reforzó el flujo de ingreso de personal con contrato de preparación, fechas de contrato y control de contrato firmado pendiente.", "Evita activar personal sin contrato firmado y ordena el paso desde ficha/ingreso hacia trabajador activo.", "En avance"),
        ("5", "Contratos", "Se ajustó la conciliación de estado laboral desde contratos y la reutilización/actualización de contratos equivalentes en preparación.", "Protege la regla de no activar trabajadores sin contrato vigente firmado y evita duplicidades innecesarias.", "En avance"),
        ("6", "Interfaz y navegación", "Se añadieron estilos y accesos para EPP/logística, mejoras responsive, menú lateral y ajustes visuales en pantallas operativas.", "Facilita uso por personal de oficina y mejora visualización en diferentes equipos.", "En avance"),
        ("7", "Pruebas", "Se actualizaron pruebas de herramientas, RQ Mina, RQ Proserge, ingresos, contratos y lista negra/disponibilidad.", "Aporta validación técnica para cambios sensibles del sistema.", "Actualizado"),
    ]
    for row, values in zip(table.rows[1:], rows):
        fill_row(row, values, alignments={0: WD_ALIGN_PARAGRAPH.CENTER, 4: WD_ALIGN_PARAGRAPH.CENTER}, size=7.25)


def add_support_summary(doc):
    heading(doc, "5. RESUMEN OPERATIVO POR TIPO DE SOPORTE")
    bullet(doc, "Infraestructura y nube: seguimiento de Nextcloud para sostener acceso y disponibilidad de información compartida.")
    bullet(doc, "Comunicaciones: configuración de correos en celulares y computadoras para mejorar continuidad de trabajo administrativo.")
    bullet(doc, "Ofimática: soporte en Excel y Office para resolver necesidades diarias de RR.HH. y oficina.")
    bullet(doc, "Equipos y periféricos: revisión de impresoras y migración de información entre discos duros.")
    bullet(doc, "Sistema Proserge: avance de logística y mejoras relacionadas con herramientas, EPP, transporte, contratos e ingresos de personal.")


def add_pending(doc):
    heading(doc, "6. PENDIENTES Y RECOMENDACIONES")
    table = fixed_table(doc, 6, 4, widths=[3.2, 5.1, 4.0, 2.3], header=True)
    fill_row(table.rows[0], ["Pendiente", "Recomendación", "Responsable sugerido", "Prioridad"], header=True, size=8.0)
    rows = [
        ("Validación de logística", "Probar con casos reales los flujos de EPP, herramientas, transporte, entregas y recepción.", "Logística / Operaciones", "Alta"),
        ("Nube Nextcloud", "Mantener control de capacidad, usuarios, permisos y respaldo de carpetas críticas.", "Sistemas", "Media"),
        ("Correos", "Documentar usuarios configurados y equipos pendientes para evitar reprocesos.", "Sistemas / Administración", "Media"),
        ("Office", "Regularizar y documentar activaciones/licenciamiento de software de oficina.", "Administración / Sistemas", "Alta"),
        ("Respaldos", "Verificar integridad de datos luego del traspaso de información entre discos.", "Sistemas", "Alta"),
    ]
    for row, values in zip(table.rows[1:], rows):
        fill_row(row, values, alignments={3: WD_ALIGN_PARAGRAPH.CENTER}, size=7.8)


def add_closing(doc):
    heading(doc, "7. CONCLUSIÓN")
    paragraph(
        doc,
        "Durante el periodo informado se realizaron labores mixtas de soporte técnico, continuidad operativa, asistencia a usuarios y avance del sistema Proserge. El trabajo combinó atención directa a oficina, soporte de herramientas digitales, seguimiento de nube y mejoras funcionales del sistema, especialmente en el frente de logística.",
    )
    paragraph(
        doc,
        "El avance de logística dentro del sistema requiere validación con usuarios reales para cerrar criterios de uso, costos, entregas, recepción y trazabilidad operativa.",
    )


def patch_header_image_alt(source_path: Path, target_path: Path):
    ns = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
    with zipfile.ZipFile(source_path, "r") as zin, zipfile.ZipFile(target_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                root = ET.fromstring(data)
                for doc_pr in root.findall(".//wp:docPr", ns):
                    doc_pr.set("title", "Membrete CDPERU")
                    doc_pr.set("descr", "Membrete institucional de CDPERU con datos de contacto.")
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(item, data)


doc = Document(TEMPLATE)
clear_body(doc)

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(5.0)
section.bottom_margin = Cm(5.0)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.3)
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

paragraph(doc, "Informe Mensual de Labores", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
paragraph(doc, "INFORME MENSUAL DE LABORES - JULIO 2026", size=15.2, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
paragraph(doc, "Soporte TI, continuidad operativa y avances del sistema PROSERGE", size=10.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=13)

heading(doc, "1. IDENTIFICACIÓN DEL INFORME")
add_identification(doc)

heading(doc, "2. RESUMEN EJECUTIVO")
paragraph(
    doc,
    "El presente informe resume las labores realizadas y reportadas para el periodo del 01 al 30 de julio de 2026. Las actividades se concentraron en reuniones de coordinación del sistema, seguimiento de la nube Nextcloud, configuración de correos, soporte ofimático, revisión de impresoras, traspaso de información entre discos y avance funcional del sistema Proserge, especialmente en logística.",
)
paragraph(
    doc,
    "En el sistema Proserge se identifican avances técnicos vinculados a logística EPP, herramientas por parada, transporte de RQ Mina, ingresos de personal, contratos, interfaz y pruebas. Estos cambios deben validarse con usuarios reales para confirmar que responden al flujo operativo de la empresa.",
)

add_labor_table(doc)
add_system_changes(doc)
add_support_summary(doc)
add_pending(doc)
add_closing(doc)

doc.core_properties.title = "Informe mensual de labores - Julio 2026"
doc.core_properties.subject = "Soporte TI y avances del sistema PROSERGE"
doc.core_properties.author = "CDPERU / Equipo de Desarrollo PROSERGE"
doc.core_properties.comments = "Documento mensual de labores y avances preparado para PROSERGE."

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(RAW_OUTPUT)
patch_header_image_alt(RAW_OUTPUT, OUTPUT)
try:
    RAW_OUTPUT.unlink()
except OSError:
    pass
print(OUTPUT)
